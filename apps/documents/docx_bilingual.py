import os
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from difflib import SequenceMatcher


def _get_word_diff(source: str, translation: str) -> list:
    """Compare words and return list of (word, is_different) tuples."""
    source_words = source.split()
    translation_words = translation.split()
    matcher = SequenceMatcher(None, source_words, translation_words)
    result = []
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'equal':
            for word in source_words[i1:i2]:
                result.append((word, False))
        elif op == 'replace':
            for word in source_words[i1:i2]:
                result.append((word, True))
            for word in translation_words[j1:j2]:
                result.append((word, True))
        elif op == 'insert':
            for word in translation_words[j1:j2]:
                result.append((word, True))
        elif op == 'delete':
            for word in source_words[i1:i2]:
                result.append((word, True))
    return result


def create_bilingual_docx(
    input_path: str,
    output_path: str,
    translate_fn,
    source_lang: str = 'en',
    target_lang: str = 'hi',
) -> dict:
    doc = Document(input_path)
    stats = {"paragraphs": 0, "characters": 0}

    new_doc = Document()

    style = new_doc.styles['Normal']
    style.font.size = Pt(11)

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            new_doc.part.rels[rel.rId] = rel

    for i, para in enumerate(doc.paragraphs):
        has_images = para._element.xpath(".//w:drawing | .//w:pict")

        if has_images:
            new_para = new_doc.add_paragraph()
            for drawing in para._element.xpath(".//w:drawing | .//w:pict"):
                new_drawing = copy.deepcopy(drawing)
                new_para._element.append(new_drawing)
            continue

        if not para.text.strip():
            new_doc.add_paragraph('')
            continue

        full_text = para.text.strip()
        if not full_text or len(full_text) < 2:
            continue

        try:
            translated = translate_fn(full_text, source_lang, target_lang)
        except Exception:
            translated = full_text

        if not translated:
            translated = full_text

        source_para = new_doc.add_paragraph()
        source_run = source_para.add_run(f"▶ {full_text}")
        source_run.font.size = Pt(10)
        source_run.font.color.rgb = RGBColor(100, 100, 100)

        target_para = new_doc.add_paragraph()
        target_para.add_run("  ")

        diff_words = _get_word_diff(full_text, translated)
        for word, is_different in diff_words:
            run = target_para.add_run(f"{word} ")
            run.font.size = Pt(11)
            if is_different:
                run.font.color.rgb = RGBColor(0, 128, 0)
                run.bold = True

        stats["paragraphs"] += 1
        stats["characters"] += len(full_text)

    for table in doc.tables:
        new_table = new_doc.add_table(rows=0, cols=len(table.columns))
        new_table.style = 'Table Grid'

        for row in table.rows:
            src_row = new_table.add_row()
            tgt_row = new_table.add_row()

            for i, cell in enumerate(row.cells):
                cell_has_images = cell._element.xpath(".//w:drawing | .//w:pict")
                cell_text = cell.text.strip()

                if cell_has_images:
                    tgt_cell = tgt_row.cells[i]
                    for p in cell.paragraphs:
                        for drawing in p._element.xpath(".//w:drawing | .//w:pict"):
                            new_drawing = copy.deepcopy(drawing)
                            tgt_cell.paragraphs[0]._element.append(new_drawing)
                    continue

                if cell_text:
                    try:
                        translated = translate_fn(cell_text, source_lang, target_lang)
                    except Exception:
                        translated = cell_text

                    src_cell = src_row.cells[i]
                    src_cell.text = f"▶ {cell_text}"
                    for p in src_cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(100, 100, 100)

                    tgt_cell = tgt_row.cells[i]
                    tgt_cell.paragraphs[0].clear()
                    tgt_cell.paragraphs[0].add_run("  ")
                    diff_words = _get_word_diff(cell_text, translated)
                    for word, is_different in diff_words:
                        run = tgt_cell.paragraphs[0].add_run(f"{word} ")
                        run.font.size = Pt(10)
                        if is_different:
                            run.font.color.rgb = RGBColor(0, 128, 0)
                            run.bold = True

        stats["paragraphs"] += 1

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    new_doc.save(output_path)
    return stats
