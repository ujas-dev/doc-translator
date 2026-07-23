import os
import copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def translate_docx_preserve(
    input_path: str,
    output_path: str,
    translate_fn,
    source_lang: str = 'en',
    target_lang: str = 'hi',
) -> dict:
    doc = Document(input_path)
    stats = {"paragraphs": 0, "tables": 0, "characters": 0}

    for para in doc.paragraphs:
        result = _translate_paragraph(para, translate_fn, source_lang, target_lang)
        if result > 0:
            stats["paragraphs"] += 1
            stats["characters"] += result

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    result = _translate_paragraph(para, translate_fn, source_lang, target_lang)
                    if result > 0:
                        stats["tables"] += 1
                        stats["characters"] += result

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return stats


def _translate_paragraph(para, translate_fn, source_lang, target_lang) -> int:
    if not para.text.strip():
        return 0

    has_page_break = para._element.xpath(".//w:br[@w:type='page'] | .//w:lastRenderedPageBreak")
    has_images = para._element.xpath(".//w:drawing | .//w:pict")
    if has_images:
        return 0

    full_text = para.text.strip()
    if not full_text or len(full_text) < 2:
        return 0

    runs_data = []
    for run in para.runs:
        if run.text.strip():
            runs_data.append({
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size,
                "font_color": run.font.color.rgb if run.font.color and run.font.color.rgb else None,
            })

    if not runs_data:
        return 0

    try:
        translated = translate_fn(full_text, source_lang, target_lang)
        if not translated or translated == full_text:
            return 0
    except Exception:
        return 0

    pPr = para._element.get_or_add_pPr()

    for run in list(para.runs):
        para._element.remove(run._element)

    if has_page_break:
        br = para._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
        run = para.add_run()
        run._element.append(br)

    if len(runs_data) == 1:
        run = para.add_run(translated)
        _apply_formatting(run, runs_data[0])
    else:
        segments = _distribute_translation(full_text, translated, runs_data)
        for seg_text, fmt in segments:
            if seg_text.strip():
                run = para.add_run(seg_text)
                _apply_formatting(run, fmt)

    return len(full_text)


def _apply_formatting(run, fmt: dict):
    if fmt.get("bold"):
        run.bold = True
    if fmt.get("italic"):
        run.italic = True
    if fmt.get("underline"):
        run.underline = True

    if fmt.get("font_name"):
        run.font.name = fmt["font_name"]
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), fmt["font_name"])
        rFonts.set(qn("w:hAnsi"), fmt["font_name"])
        rFonts.set(qn("w:eastAsia"), fmt["font_name"])

    if fmt.get("font_size"):
        run.font.size = fmt["font_size"]

    if fmt.get("font_color"):
        run.font.color.rgb = fmt["font_color"]


def _distribute_translation(original: str, translated: str, runs_data: list) -> list:
    if not runs_data:
        return [(translated, {})]

    total_len = sum(len(r["text"]) for r in runs_data)
    if total_len == 0:
        return [(translated, runs_data[0])]

    result = []
    chars_used = 0
    for i, rd in enumerate(runs_data):
        proportion = len(rd["text"]) / total_len
        seg_len = int(len(translated) * proportion)

        if i == len(runs_data) - 1:
            seg_text = translated[chars_used:]
        else:
            seg_text = translated[chars_used:chars_used + seg_len]
            chars_used += seg_len

        result.append((seg_text, rd))

    return result
