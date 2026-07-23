import pytest


class TestDocxLayoutModule:
    def test_import_write_docx_layout(self):
        from apps.documents.parsers import write_docx_layout
        assert callable(write_docx_layout)

    def test_write_docx_layout_returns_stats(self, tmp_path):
        from docx import Document
        from apps.documents.parsers import write_docx_layout

        input_path = str(tmp_path / "input.docx")
        output_path = str(tmp_path / "output.docx")

        doc = Document()
        doc.add_paragraph("Hello world")
        doc.add_paragraph("This is a test document.")
        doc.save(input_path)

        def translate_fn(text, source, target):
            return f"Translated: {text}"

        stats = write_docx_layout(
            input_path=input_path,
            output_path=output_path,
            translate_fn=translate_fn,
            source_lang="en",
            target_lang="hi",
        )
        assert 'paragraphs' in stats
        assert stats['paragraphs'] >= 2

    def test_write_docx_layout_preserves_structure(self, tmp_path):
        from docx import Document
        from apps.documents.parsers import write_docx_layout

        input_path = str(tmp_path / "input.docx")
        output_path = str(tmp_path / "output.docx")

        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Body text")
        doc.save(input_path)

        def translate_fn(text, source, target):
            return f"Translated: {text}"

        write_docx_layout(
            input_path=input_path,
            output_path=output_path,
            translate_fn=translate_fn,
            source_lang="en",
            target_lang="hi",
        )

        output_doc = Document(output_path)
        assert len(output_doc.paragraphs) >= 2
        texts = [p.text for p in output_doc.paragraphs]
        assert any("Title" in t or "Translated" in t for t in texts)
