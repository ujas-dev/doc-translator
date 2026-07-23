import os
import tempfile
import pytest
from unittest.mock import MagicMock


class TestPdfBilingual:
    def test_import(self):
        from apps.documents.pdf_bilingual import create_bilingual_pdf
        assert callable(create_bilingual_pdf)

    def test_bilingual_pdf_creation(self):
        import fitz
        from apps.documents.pdf_bilingual import create_bilingual_pdf

        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "input.pdf")
            output_path = os.path.join(tmpdir, "output.pdf")

            doc = fitz.open()
            page = doc.new_page()
            page.insert_text(fitz.Point(50, 100), "Hello world", fontsize=12)
            doc.save(input_path)
            doc.close()

            def translate_fn(text, source, target):
                return f"Translated: {text}"

            stats = create_bilingual_pdf(
                input_path=input_path,
                output_path=output_path,
                translate_fn=translate_fn,
                source_lang="en",
                target_lang="hi",
            )

            assert os.path.exists(output_path)
            assert stats["pages"] == 1
            assert stats["segments"] >= 0


class TestDocxBilingual:
    def test_import(self):
        from apps.documents.docx_bilingual import create_bilingual_docx
        assert callable(create_bilingual_docx)

    def test_bilingual_docx_creation(self):
        from docx import Document
        from apps.documents.docx_bilingual import create_bilingual_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, "input.docx")
            output_path = os.path.join(tmpdir, "output.docx")

            doc = Document()
            doc.add_paragraph("Hello world")
            doc.save(input_path)

            def translate_fn(text, source, target):
                return f"Translated: {text}"

            stats = create_bilingual_docx(
                input_path=input_path,
                output_path=output_path,
                translate_fn=translate_fn,
                source_lang="en",
                target_lang="hi",
            )

            assert os.path.exists(output_path)
            assert stats["paragraphs"] >= 1

            output_doc = Document(output_path)
            paragraphs = [p.text for p in output_doc.paragraphs]
            assert any("Hello world" in p for p in paragraphs)
            assert any("Translated:" in p for p in paragraphs)


class TestBilingualModelField:
    def test_bilingual_default_false(self):
        from apps.documents.models import DocumentJob
        field = DocumentJob._meta.get_field('bilingual')
        assert field.default is False

    def test_bilingual_field_type(self):
        from apps.documents.models import DocumentJob
        from django.db import models as django_models
        field = DocumentJob._meta.get_field('bilingual')
        assert isinstance(field, django_models.BooleanField)
